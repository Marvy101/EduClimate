import streamlit as st
from streamlit_tags import st_tags
from docx import Document
from io import BytesIO
import json

# Set the title of your application
st.title("Generate a Lesson Plan for Climate related issues!")

# Create a form
with st.form(key='my_form'):
    # Create a text input for the lesson title
    lesson_title = st.text_input("Enter Lesson Title")

    all_topics = st_tags(
        label='Enter Key Topics:',
        text='Write topics and press enter to add',
        value=[],
        suggestions=['CO2', 'Keeling Curve', 'Solar Panels'],
        maxtags=10,
        key='4'
    )

    learning_objectives = st_tags(
        label='Enter learning_objectives :',
        text='Write objectives and press enter to add',
        value=[],
        suggestions=[],
        maxtags=10,
        key='5'
    )

    st.markdown("# Audience")
    all_audience = st_tags(
        label='Enter audiences:',
        text='Write audiences and press enter to add',
        value=[],
        suggestions=["High School", "General Audience", "AP"],
        maxtags=10,
        key='6'
    )

    # Create a text input for number of students
    number_of_students = st.text_input("How many students?")

    all_locations = st_tags(
        label='Enter locations:',
        text='Write locations and press enter to add',
        value=[],
        suggestions=["classroom", "field", "Bay Area"],
        maxtags=10,
        key='7'
    )

    all_acrqs = st_tags(
        label='Enter Accessibilty Requirments:',
        text='Write Accessibilty Requirments and press enter to add',
        value=[],
        suggestions=["Students with Visual Impairements"],
        maxtags=10,
        key='8'
    )

    st.markdown("# Classroom Practicalities")
    all_avrs = st_tags(
        label='Enter Available Resources:',
        text='Write Available Resources and press enter to add',
        value=[],
        suggestions=["WhiteBoard","Computers"],
        maxtags=10,
        key='9'
    )

    # Create a text input for lesson length
    lesson_length = st.text_input("Lesson Length (minutes): ")

    all_teme = st_tags(
        label='Enter Teaching Methods:',
        text='Write Teaching Methods and press enter to add',
        value=[],
        suggestions=["Lecture", "Discussion", "Hands-On"],
        maxtags=10,
        key='10'
    )
    # Create a submit button inside the form
    submit_button = st.form_submit_button(label='Generate')

# You can use the variables from the form anywhere in your code
if submit_button:

    # Print all the variables
    st.write("Lesson Title: ", lesson_title)
    st.write("All Topics: ", all_topics)
    st.write("Learning Objectives: ", learning_objectives)
    st.write("All Audience: ", all_audience)
    st.write("Number of Students: ", number_of_students)
    st.write("All Locations: ", all_locations)
    st.write("All Accessibility Requirements: ", all_acrqs)
    st.write("All Available Resources: ", all_avrs)
    st.write("Lesson Length: ", lesson_length)
    st.write("All Teaching Methods: ", all_teme)

    final = [lesson_title, all_topics, learning_objectives, all_audience, number_of_students, all_locations, all_acrqs, all_avrs, lesson_length, all_teme ]
    print(final)

    #we can call for the api here using final

    json_data = json.dumps(final)

    # Convert the JSON string to a Python list of dictionaries
    data = json.loads(json_data)

    # Create a new Word document
    doc = Document()

    # We can format this to have titles in a docx file
    doc.add_paragraph(str(data))

    # Save the Word document to a BytesIO object
    b = BytesIO()
    doc.save(b)
    b.seek(0)

    # Create a download link for the Word document
    st.download_button(
        label="Download Word document",
        data=b,
        file_name='document.docx',
        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )

json_data = """
{
  "introduction": "Begin the lesson with a brief discussion of climate change, highlighting the role of carbon dioxide as a greenhouse gas. Explain how the increase in greenhouse gases, particularly CO2, leads to global warming and ultimately climate change. Introduce the concept of the Keeling Curve, its significance, and its importance in understanding the relationship between atmospheric CO2 concentrations and global climate change.",
  "core_content": "Explain the main contributors of CO2 emissions, such as fossil fuel burning (coal, oil, and natural gas), deforestation, and other land-use changes. Discuss how these human activities have led to increased atmospheric CO2 concentrations over time. Present the data from the Keeling Curve, showing the increasing trend of CO2 concentrations from the late 1950s to current times. Explain the seasonality seen in the curve and its relation to plant growth and decay.",
  "learning_activities": [
    "1. Distribute the Keeling Curve handout to students, either individually or in groups. Have them carefully examine the graph and annotate it, noting significant events (e.g., 400 ppm milestone) and relating these events to human activities. Encourage students to also observe the seasonal variations in the curve.",
    "2. Using the trend shown on the Keeling Curve, encourage students to predict likely future increases in CO2 levels. Students can create their own rough graph projections, or use a simple linear equation to calculate estimated future CO2 concentrations. Discuss the potential consequences if CO2 emissions continue to increase.",
    "3. Discuss the potential impacts of increasing atmospheric CO2 levels on global climate patterns, ecosystems, and human society. Allow students to share their thoughts in a class discussion or small group discussions. Encourage them to consider both physical consequences (e.g., sea level rise, extreme weather events) and societal consequences (e.g., food security, migration)."
  ],
  "assessment_evaluation": "Assess students' understanding of the relationship between atmospheric carbon dioxide and climate change through their predictions and discussions. Consider evaluating their ability to recognize trends in the Keeling Curve data, the accuracy of their predictions based on that data, and their understanding of the potential impacts of continuing CO2 emissions on the climate system and human societies.",
  "conclusion": "Summarize the findings of the class, highlighting the main points of discussion and the potential implications of continuing CO2 emissions on our planet. Encourage students to consider personal and societal actions that could help reduce CO2 emissions and mitigate the potential impacts of climate change, such as adopting renewable energy sources, reforestation efforts, and reducing personal carbon footprints."
}
"""

# Convert the JSON string to a Python dictionary
data = json.loads(json_data)

# Create a new Word document
doc = Document()

# Add a section to the Word document for each item in the data
for key, value in data.items():
    doc.add_heading(key.replace('_', ' ').title(), level=1)
    if isinstance(value, list):
        for item in value:
            doc.add_paragraph(item)
    else:
        doc.add_paragraph(value)

# Save the Word document to a BytesIO object
b = BytesIO()
doc.save(b)
b.seek(0)

# Create a download link for the Word document
st.download_button(
    label="Download Word document",
    data=b,
    file_name='document.docx',
    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
)